from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document
import io

app = Flask(__name__)
CORS(app)

@app.route('/parse-word', methods=['POST'])
def parse_word():
    """Reçoit un fichier Word .docx et retourne les données extraites sous forme de JSON"""
    if 'file' not in request.files:
        return jsonify({'error': 'Aucun fichier reçu'}), 400

    uploaded_file = request.files['file']
    if uploaded_file.filename == '':
        return jsonify({'error': 'Fichier vide ou non sélectionné'}), 400

    try:
        doc_bytes = uploaded_file.read()
        document = Document(io.BytesIO(doc_bytes))
        emplois = []

        if document.tables:
            table = document.tables[0]
            rows = table.rows

            for row in rows[1:]:  # Ignorer l'en-tête
                cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                if any(cells):  # Ignore les lignes totalement vides
                    entry = {}
                    keys = ['classe', 'jour', 'heure', 'module', 'prof', 'salle']
                    for idx, key in enumerate(keys):
                        if idx < len(cells):
                            entry[key] = cells[idx]
                    emplois.append(entry)

        else:
            for para in document.paragraphs:
                line = para.text.strip()
                if not line:
                    continue
                parts = [p.strip() for p in line.split('|')]
                if len(parts) >= 6:
                    entry = dict(zip(['classe', 'jour', 'heure', 'module', 'prof', 'salle'], parts[:6]))
                    emplois.append(entry)

        return jsonify({'emplois': emplois})

    except Exception as e:
        return jsonify({'error': f'Erreur lors du traitement : {str(e)}'}), 500

if __name__ == '__main__':
    # À utiliser uniquement pour des tests
    app.run(host='0.0.0.0', port=8000)
