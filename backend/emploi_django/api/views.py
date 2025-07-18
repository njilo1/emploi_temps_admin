from rest_framework import viewsets, status
from rest_framework.decorators import api_view
from rest_framework.response import Response
from django.http import JsonResponse
from docx import Document
import io
import unicodedata
from .models import Classe, Filiere, Salle, Module, Professeur, Emploi, Departement
from .serializers import (
    ClasseSerializer, FiliereSerializer, SalleSerializer,
    ModuleSerializer, ProfesseurSerializer, EmploiSerializer, DepartementSerializer
)

def clean_text(text):
    """Nettoie le texte en remplaçant les caractères spéciaux"""
    if not text:
        return text
    
    # Remplacer les caractères spéciaux courants
    replacements = {
        'é': 'e', 'è': 'e', 'ê': 'e', 'ë': 'e',
        'à': 'a', 'â': 'a', 'ä': 'a',
        'î': 'i', 'ï': 'i',
        'ô': 'o', 'ö': 'o',
        'ù': 'u', 'û': 'u', 'ü': 'u',
        'ç': 'c',
        '–': '-', '—': '-',  # Tirets
        "'": "'", '"': '"',  # Guillemets
    }
    
    result = text
    for old, new in replacements.items():
        result = result.replace(old, new)
    
    return result

# ======== CRUD ViewSets ========
class ClasseViewSet(viewsets.ModelViewSet):
    queryset = Classe.objects.all()
    serializer_class = ClasseSerializer

class FiliereViewSet(viewsets.ModelViewSet):
    queryset = Filiere.objects.all()
    serializer_class = FiliereSerializer

class DepartementViewSet(viewsets.ModelViewSet):
    queryset = Departement.objects.all()
    serializer_class = DepartementSerializer

class SalleViewSet(viewsets.ModelViewSet):
    queryset = Salle.objects.all()
    serializer_class = SalleSerializer

class ModuleViewSet(viewsets.ModelViewSet):
    queryset = Module.objects.all()
    serializer_class = ModuleSerializer

class ProfesseurViewSet(viewsets.ModelViewSet):
    queryset = Professeur.objects.all()
    serializer_class = ProfesseurSerializer

class EmploiViewSet(viewsets.ModelViewSet):
    queryset = Emploi.objects.all()
    serializer_class = EmploiSerializer

# ======== Route POST /emplois/generate/ ========
@api_view(['POST'])
def generer_emplois(request):
    try:
        tranches_horaires = [
            '07H30 - 10H00',
            '10H15 - 12H45',
            '13H00 - 15H30',
            '15H45 - 18H15',
        ]
        jours_semaine = ['Lundi', 'Mardi', 'Mercredi', 'Jeudi', 'Vendredi', 'Samedi']

        Emploi.objects.all().delete()

        classes = Classe.objects.all()
        salles = Salle.objects.filter(disponible=True)
        modules = Module.objects.all()

        salle_occupe = {}
        prof_occupe = {}

        for classe in classes:
            for module in modules:
                heures_restantes = getattr(module, 'volume_horaire', 3)
                prof = getattr(module, 'prof', None)
                if not prof:
                    continue

                for jour in jours_semaine:
                    for heure in tranches_horaires:
                        cle = f"{jour}-{heure}"
                        salle_id = None

                        for salle in salles:
                            if salle.capacite >= classe.effectif and cle not in salle_occupe.get(salle.id, set()):
                                salle_id = salle.id
                                salle_occupe.setdefault(salle.id, set()).add(cle)
                                break

                        prof_id = prof.id
                        if salle_id and cle not in prof_occupe.get(prof_id, set()):
                            prof_occupe.setdefault(prof_id, set()).add(cle)

                            Emploi.objects.create(
                                classe=classe,
                                module=module,
                                prof=prof,
                                salle=Salle.objects.get(id=salle_id),
                                jour=jour,
                                heure=heure
                            )
                            heures_restantes -= 3
                            if heures_restantes <= 0:
                                break
                    if heures_restantes <= 0:
                        break

        return Response({"message": "Emplois générés avec succès"}, status=200)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return Response({"error": str(e)}, status=500)

# ======== Route GET /emplois/classe/<id>/ ========
@api_view(['GET'])
def emploi_par_classe(request, classe_id):
    try:
        classe = Classe.objects.get(id=classe_id)
        emplois = Emploi.objects.filter(classe=classe)

        print(f"🔍 Récupération emplois pour classe {classe.nom} (ID: {classe_id})")
        print(f"📊 Nombre d'emplois trouvés: {emplois.count()}")

        data = {}
        for emploi in emplois:
            jour = emploi.jour
            heure = emploi.heure
            # Nettoyer les caractères spéciaux
            module_nom = clean_text(emploi.module.nom)
            salle_nom = clean_text(emploi.salle.nom)
            prof_nom = clean_text(emploi.prof.nom)
            
            # Format avec sauts de ligne pour un affichage plus propre
            libelle = f"{module_nom}\n{salle_nom}\n{prof_nom}"

            if jour not in data:
                data[jour] = {}
            data[jour][heure] = libelle
            
            print(f"📅 Ajout: {jour} {heure} -> {libelle}")

        print(f"📤 Données retournées: {data}")
        return Response(data, content_type='application/json; charset=utf-8')
    except Classe.DoesNotExist:
        print(f"❌ Classe introuvable: {classe_id}")
        return Response({"error": "Classe introuvable"}, status=404)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return Response({"error": str(e)}, status=500)

# ======== Route POST /emplois/import/ ========
@api_view(['POST'])
def import_emplois(request):
    try:
        import json
        print(f"📦 Requête brute: {request.body}")
        try:
            body = json.loads(request.body)
            print(f"📦 Body JSON: {body}")
        except Exception as e:
            print(f"❌ Erreur parsing JSON: {e}")
        # Gérer les deux formats possibles : request.data ou request.data.get('emplois')
        if isinstance(request.data, list):
            data = request.data
        else:
            data = request.data.get('emplois', [])
        
        print(f"📥 Données reçues: {data}")  # Debug log
        
        if not data:
            return Response({"error": "Aucune donnée à importer"}, status=400)

        Emploi.objects.all().delete()
        emplois_crees = 0

        for item in data:
            try:
                # Gérer les deux formats : noms ou IDs
                if isinstance(item.get('classe'), int):
                    # Format avec IDs
                    classe_id = item.get('classe')
                    module_id = item.get('module')
                    prof_id = item.get('prof')
                    salle_id = item.get('salle')
                    jour = item.get('jour', '').strip()
                    heure = item.get('heure', '').strip()

                    try:
                        classe = Classe.objects.get(id=classe_id)
                        module = Module.objects.get(id=module_id)
                        prof = Professeur.objects.get(id=prof_id)
                        salle = Salle.objects.get(id=salle_id)
                    except (Classe.DoesNotExist, Module.DoesNotExist, Professeur.DoesNotExist, Salle.DoesNotExist):
                        print(f"⚠️ ID introuvable ignoré: {item}")
                        continue

                else:
                    # Format avec noms (nouveau format)
                    classe_nom = item.get('classe', '').strip()
                    module_nom = item.get('module', '').strip()
                    prof_nom = item.get('prof', '').strip()
                    salle_nom = item.get('salle', '').strip()
                    jour = item.get('jour', '').strip()
                    heure = item.get('heure', '').strip()

                    print(f"🔍 Traitement de l'item: {item}")  # Debug log

                    if not (classe_nom and module_nom and prof_nom and salle_nom and jour and heure):
                        print(f"⚠️ Donnée incomplète ignorée: {item}")
                        continue

                    # Créer une filière par défaut si nécessaire
                    filiere, _ = Filiere.objects.get_or_create(nom="Générale")
                    
                    classe, _ = Classe.objects.get_or_create(
                        nom=clean_text(classe_nom), 
                        defaults={"effectif": 30, "filiere": filiere}
                    )
                    prof, _ = Professeur.objects.get_or_create(nom=clean_text(prof_nom))
                    salle, _ = Salle.objects.get_or_create(nom=clean_text(salle_nom), defaults={"capacite": 30, "disponible": True})
                    module, _ = Module.objects.get_or_create(nom=clean_text(module_nom), defaults={"classe": classe, "prof": prof})

                    if module.classe != classe or module.prof != prof:
                        module.classe = classe
                        module.prof = prof
                        module.save()

                # Valider que le jour est dans les choix autorisés
                jours_valides = ['Lundi', 'Mardi', 'Mercredi', 'Jeudi', 'Vendredi', 'Samedi']
                if jour not in jours_valides:
                    print(f"⚠️ Jour invalide ignoré: {jour} pour l'item {item}")
                    continue

                Emploi.objects.create(
                    classe=classe,
                    module=module,
                    prof=prof,
                    salle=salle,
                    jour=jour,
                    heure=heure
                )
                emplois_crees += 1
                
            except Exception as item_error:
                print(f"❌ Erreur lors du traitement de l'item {item}: {item_error}")
                continue

        return Response({
            "message": f"Importation réussie: {emplois_crees} emplois créés",
            "emplois_crees": emplois_crees
        }, status=200)

    except Exception as e:
        import traceback
        traceback.print_exc()
        return Response({"error": str(e)}, status=500)

# ======== Route DELETE /emplois/clear/ ========
@api_view(['DELETE'])
def clear_emplois(request):
    try:
        count = Emploi.objects.count()
        Emploi.objects.all().delete()
        return Response({
            "message": f"Base de données vidée : {count} emplois supprimés"
        }, status=200)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return Response({"error": str(e)}, status=500)

# ======== Route POST /api/parse-word/ ========
@api_view(['POST'])
def parse_word_file(request):
    """Reçoit un fichier Word .docx et retourne les données extraites sous forme de JSON"""
    try:
        if 'file' not in request.FILES:
            return Response({'error': 'Aucun fichier reçu'}, status=400)

        uploaded_file = request.FILES['file']
        if uploaded_file.name == '':
            return Response({'error': 'Fichier vide ou non sélectionné'}, status=400)

        doc_bytes = uploaded_file.read()
        document = Document(io.BytesIO(doc_bytes))
        emplois = []

        if document.tables:
            table = document.tables[0]
            rows = table.rows

            for row in rows[1:]:  # Ignorer l'en-tête
                cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                if any(cells):  # Ignore les lignes totalement vides
                    entry = {}
                    keys = ['classe', 'jour', 'heure', 'module', 'prof', 'salle']
                    for idx, key in enumerate(keys):
                        if idx < len(cells):
                            entry[key] = cells[idx]
                    emplois.append(entry)

        else:
            for para in document.paragraphs:
                line = para.text.strip()
                if not line:
                    continue
                parts = [p.strip() for p in line.split('|')]
                if len(parts) >= 6:
                    entry = dict(zip(['classe', 'jour', 'heure', 'module', 'prof', 'salle'], parts[:6]))
                    emplois.append(entry)

        # Sauvegarder dans la base de données Django
        for emploi_data in emplois:
            classe_nom = emploi_data.get('classe', '')
            module_nom = emploi_data.get('module', '')
            prof_nom = emploi_data.get('prof', '')
            salle_nom = emploi_data.get('salle', '')
            jour = emploi_data.get('jour', '')
            heure = emploi_data.get('heure', '')

            if classe_nom and module_nom and prof_nom and salle_nom and jour and heure:
                # Créer ou récupérer les objets
                classe, _ = Classe.objects.get_or_create(nom=classe_nom, defaults={"effectif": 30})
                prof, _ = Professeur.objects.get_or_create(nom=prof_nom)
                salle, _ = Salle.objects.get_or_create(nom=salle_nom, defaults={"capacite": 30, "disponible": True})
                module, _ = Module.objects.get_or_create(nom=module_nom, defaults={"classe": classe, "prof": prof})

                # Mettre à jour les relations si nécessaire
                if module.classe != classe or module.prof != prof:
                    module.classe = classe
                    module.prof = prof
                    module.save()

                # Créer l'emploi du temps
                Emploi.objects.create(
                    classe=classe,
                    module=module,
                    prof=prof,
                    salle=salle,
                    jour=jour,
                    heure=heure
                )

        return Response({
            'emplois': emplois, 
            'message': 'Données extraites et sauvegardées avec succès'
        }, status=200)

    except Exception as e:
        import traceback
        traceback.print_exc()
        return Response({'error': f'Erreur lors du traitement : {str(e)}'}, status=500)
