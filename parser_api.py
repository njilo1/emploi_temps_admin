from fastapi import FastAPI, File, UploadFile
from fastapi.responses import JSONResponse
from docx import Document
from typing import Dict
import uvicorn

app = FastAPI()


def parse_docx(path: str) -> Dict[str, Dict[str, Dict[str, str]]]:
    """Very basic parser returning data grouped by class."""
    doc = Document(path)
    result: Dict[str, Dict[str, Dict[str, str]]] = {}

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 4:
                classe = cells[0].replace(' ', '_')
                jour = cells[1]
                heure = cells[2]
                detail = cells[3]
                result.setdefault(classe, {}).setdefault(jour, {})[heure] = detail
    # fallback: try paragraphs formatted with commas
    if not result:
        for para in doc.paragraphs:
            parts = [p.strip() for p in para.text.split(',')]
            if len(parts) >= 4:
                classe = parts[0].replace(' ', '_')
                jour = parts[1]
                heure = parts[2]
                detail = ','.join(parts[3:]).strip()
                result.setdefault(classe, {}).setdefault(jour, {})[heure] = detail
    return result


@app.post("/parse-word")
async def parse_word(file: UploadFile = File(...)):
    data = await file.read()
    tmp_path = f"/tmp/{file.filename}"
    with open(tmp_path, "wb") as f:
        f.write(data)
    parsed = parse_docx(tmp_path)
    return JSONResponse(content=parsed)


if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
